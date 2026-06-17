import io
import json
import os
import uuid
import re
import pandas as pd
import streamlit as st
from dotenv import load_dotenv

try:
    from docx import Document
    from docx.shared import RGBColor
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
# Load environment variables from .env.agentbase or .env
load_dotenv(dotenv_path=".env.agentbase")
load_dotenv(dotenv_path=".env")

from agent.main import check_compliance_mock, extract_contract_risks, generate_negotiation_email, run_agent
from agent.memory import init_session, get_memory_client
from interview.question_generator import extract_text_from_file

def format_summary_html(summary_text: str, warnings: list) -> str:
    """Format contract analysis summary as HTML with Enterprise CSS styling."""
    # Bỏ lùi lề để Streamlit không nhận nhầm thành Markdown Code Block
    html_content = """<style>
.info-card {
    background-color: #f0f8ff;
    border-left: 6px solid #0078d4;
    padding: 20px;
    border-radius: 8px;
    margin-bottom: 20px;
    box-shadow: 0 2px 4px rgba(0,0,0,0.05);
}
.warning-card {
    background-color: #fff9e6;
    border-left: 6px solid #ffb900;
    padding: 20px;
    border-radius: 8px;
    margin-bottom: 20px;
    box-shadow: 0 2px 4px rgba(0,0,0,0.05);
}
.card-header {
    font-size: 1.2em;
    font-weight: 600;
    margin-bottom: 12px;
    color: #333;
    display: flex;
    align-items: center;
}
.warning-title {
    font-weight: bold;
    color: #b28200;
    background-color: #ffe082;
    padding: 3px 8px;
    border-radius: 4px;
    font-size: 0.9em;
    display: inline-block;
    margin-bottom: 6px;
}
.warning-item {
    margin-bottom: 16px;
    line-height: 1.5;
}
.warning-item:last-child {
    margin-bottom: 0;
}
</style>
"""
    
    html_content += f"""<div class="info-card">
<div class="card-header">📄 Tóm tắt chung</div>
<div style="color: #444; line-height: 1.6; font-size: 15px;">{summary_text}</div>
</div>
"""
    
    if warnings:
        html_content += """<div class="warning-card">
<div class="card-header">🟡 Điểm cần thận trọng (Cảnh báo Vàng)</div>
<div style="color: #444;">
"""
        for item in warnings:
            warn_text = item if isinstance(item, str) else item.get('warning', item.get('description', 'N/A'))
            warn_desc = "" if isinstance(item, str) else item.get('description', '')
            html_content += f"""<div class="warning-item">
<span class="warning-title">{warn_text}</span><br>
{f'<span style="font-size: 0.95em; color: #555;">{warn_desc}</span>' if warn_desc else ''}
</div>
"""
        html_content += "</div></div>"
    
    return html_content


def format_compliance_html(compliance_dict: dict) -> str:
    """Format compliance analysis JSON as structured HTML with Enterprise CSS styling."""
    html = """<style>
.comp-card {
    padding: 16px;
    border-radius: 8px;
    margin-bottom: 20px;
    box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    font-family: -apple-system, BlinkMacSystemFont, sans-serif;
}
.comp-summary { background-color: #e3f2fd; border-left: 5px solid #1976d2; }
.comp-legal { background-color: #ffebee; border-left: 5px solid #d32f2f; }
.comp-finance { background-color: #fff3e0; border-left: 5px solid #f57c00; }
.comp-nego { background-color: #e8f5e9; border-left: 5px solid #388e3c; }
.comp-header { font-size: 1.15em; font-weight: 600; margin-bottom: 12px; }
.item-title { font-weight: 600; margin-bottom: 4px; }
.item-desc { color: #555; font-size: 0.95em; line-height: 1.5; margin-bottom: 12px; padding-bottom: 12px; border-bottom: 1px dashed rgba(0,0,0,0.1); }
.item-desc:last-child { margin-bottom: 0; padding-bottom: 0; border-bottom: none; }
</style>
<div>
"""
    
    if "summary" in compliance_dict:
        html += f"""<div class="comp-card comp-summary">
<div class="comp-header" style="color: #1565c0;">📋 Tóm Tắt Thẩm định</div>
<div style="color: #424242; line-height: 1.6;">{compliance_dict['summary']}</div>
</div>
"""
    
    if "legal_risks" in compliance_dict and compliance_dict["legal_risks"]:
        html += '<div class="comp-card comp-legal"><div class="comp-header" style="color: #c62828;">⚖️ Rủi Ro Pháp Lý</div>'
        for idx, item in enumerate(compliance_dict["legal_risks"], 1):
            html += f"""<div class="item-title" style="color: #b71c1c;">#{idx}. {item.get('risk', 'N/A')}</div>
<div class="item-desc">{item.get('description', '')}</div>
"""
        html += '</div>'
    
    if "financial_risks" in compliance_dict and compliance_dict["financial_risks"]:
        html += '<div class="comp-card comp-finance"><div class="comp-header" style="color: #e65100;">💰 Rủi Ro Tài Chính</div>'
        for idx, item in enumerate(compliance_dict["financial_risks"], 1):
            html += f"""<div class="item-title" style="color: #e65100;">#{idx}. {item.get('risk', 'N/A')}</div>
<div class="item-desc">{item.get('description', '')}</div>
"""
        html += '</div>'
    
    if "negotiation_points" in compliance_dict and compliance_dict["negotiation_points"]:
        html += '<div class="comp-card comp-nego"><div class="comp-header" style="color: #2e7d32;">💬 Điểm Đề Xuất Đàm Phán</div>'
        for idx, item in enumerate(compliance_dict["negotiation_points"], 1):
            html += f"""<div class="item-title" style="color: #1b5e20;">🎯 {item.get('point', 'N/A')}</div>
<div class="item-desc">👉 <strong>Đề xuất:</strong> {item.get('suggestion', '')}</div>
"""
        html += '</div>'
    
    html += '</div>'
    return html

def generate_annotated_docx(analysis_dict: dict, original_file_bytes: bytes) -> io.BytesIO:
    """Đọc file DOCX gốc, highlight các đoạn rủi ro và chèn comment của AI trực tiếp vào dưới đoạn đó."""
    if not HAS_DOCX or not original_file_bytes:
        return None
        
    try:
        from docx.enum.text import WD_COLOR_INDEX
        from docx.shared import Pt, RGBColor
        
        file_stream = io.BytesIO(original_file_bytes)
        doc = Document(file_stream)
        
        # Lấy danh sách rủi ro và cảnh báo
        risks = []
        if "legal_df" in analysis_dict and not analysis_dict["legal_df"].empty:
            for _, row in analysis_dict["legal_df"].iterrows():
                risks.append({
                    "issue": str(row.get('Vấn đề', '')),
                    "desc": str(row.get('Rủi ro', '')),
                    "sugg": str(row.get('Đề xuất', '')),
                    "keywords": str(row.get('Nhóm điều khoản', '')).lower()
                })
                
        # Thêm các cảnh báo chung (yellow_warnings) vào danh sách cần quét
        for w in analysis_dict.get("yellow_warnings", []):
            warn_text = w if isinstance(w, str) else w.get('warning', w.get('description', ''))
            warn_desc = "" if isinstance(w, dict) else w.get('description', '')
            risks.append({
                "issue": warn_text,
                "desc": warn_desc,
                "sugg": "Cần làm rõ với đối tác.",
                "keywords": warn_text.lower()
            })

        # Bộ từ khóa mở rộng để Agent map rủi ro vào đúng đoạn văn trong file gốc
        keyword_map = {
            "thanh toán": ["thanh toán", "đợt", "hóa đơn", "tạm ứng"],
            "phạt": ["vi phạm", "phạt", "bồi thường", "thiệt hại"],
            "chấm dứt": ["chấm dứt", "hủy bỏ", "phá sản"],
            "bảo mật": ["bảo mật", "tiết lộ", "thông tin"],
            "tuân thủ": ["rửa tiền", "xung đột lợi ích", "hối lộ", "trừng phạt"],
            "thời hạn": ["thời hạn", "ngày", "hiệu lực"]
        }

        # Duyệt qua từng đoạn văn trong Hợp đồng để quét lỗi
        for p in doc.paragraphs:
            p_text = p.text.lower()
            if len(p_text) < 15:  # Bỏ qua các dòng quá ngắn (tiêu đề, dòng trống)
                continue
                
            matched_risk = None
            for risk in risks:
                is_match = False
                
                # 1. Tìm theo bộ từ khóa nhóm
                for key, kws in keyword_map.items():
                    if key in risk["keywords"] or key in risk["issue"].lower():
                        if any(kw in p_text for kw in kws):
                            is_match = True
                            break
                            
                # 2. Tìm trực tiếp cụm từ vấn đề (issue) trong văn bản
                issue_words = [w for w in risk["issue"].lower().split() if len(w) > 4]
                if issue_words and any(word in p_text for word in issue_words):
                    is_match = True
                    
                if is_match:
                    matched_risk = risk
                    break
            
            # Nếu phát hiện đoạn văn bản này có rủi ro
            if matched_risk:
                # Bôi vàng (Highlight) toàn bộ đoạn văn
                for run in p.runs:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    
                # Chèn comment của AI ngay bên dưới đoạn văn
                comment_run = p.add_run(f"\n\n[🔍 GREENNODE AI COMMENT]\n⚠️ Vấn đề: {matched_risk['issue']}\n👉 Rủi ro: {matched_risk['desc']}\n💡 Đề xuất: {matched_risk['sugg']}\n")
                comment_run.bold = True
                comment_run.font.color.rgb = RGBColor(220, 53, 69)  # Text màu đỏ tươi
                comment_run.font.size = Pt(10)
                
                # Loại bỏ rủi ro đã được comment để Agent đi quét các rủi ro khác
                risks.remove(matched_risk)

        # Chèn Báo cáo tổng quan vào trang cuối cùng của Hợp đồng
        doc.add_page_break()

        # SỬA LỖI: Dùng thẻ text tự định dạng thay vì dùng Heading 0 để tránh lỗi thiếu Style
        heading_p = doc.add_paragraph()
        heading_run = heading_p.add_run('BÁO CÁO TỔNG QUAN TỪ GREENNODE GUARDIAN')
        heading_run.bold = True
        heading_run.font.size = Pt(16)
        heading_run.font.color.rgb = RGBColor(0, 120, 212) # Màu xanh dương
        
        doc.add_paragraph(analysis_dict.get("summary", ""))

        
        #doc.add_heading('BÁO CÁO TỔNG QUAN TỪ GREENNODE GUARDIAN', 0)
       # doc.add_paragraph(analysis_dict.get("summary", ""))

        out_buffer = io.BytesIO()
        doc.save(out_buffer)
        out_buffer.seek(0)
        return out_buffer
        
    except Exception as e:
        print(f"Lỗi khi xử lý DOCX: {e}")
        return None
    

# def generate_annotated_docx(analysis_dict: dict, original_file_bytes: bytes) -> io.BytesIO:
#     """Mở file DOCX gốc của user và đính kèm Báo cáo rà soát lên đầu trang."""
#     if not HAS_DOCX or not original_file_bytes:
#         return None
        
#     try:
#         # Mở file gốc của user
#         file_stream = io.BytesIO(original_file_bytes)
#         doc = Document(file_stream)
        
#         # Tạo một tài liệu mới để ghép (Báo cáo đứng trước, Hợp đồng gốc đứng sau)
#         merged_doc = Document()
        
#         # --- PHẦN 1: BÁO CÁO RÀ SOÁT CỦA AI ---
#         merged_doc.add_heading('BÁO CÁO RÀ SOÁT HỢP ĐỒNG - GREENNODE GUARDIAN', 0)
        
#         merged_doc.add_heading('1. TÓM TẮT CHUNG', level=1)
#         merged_doc.add_paragraph(analysis_dict.get("summary", ""))
        
#         warnings = analysis_dict.get("yellow_warnings", [])
#         if warnings:
#             merged_doc.add_heading('2. CẢNH BÁO (THẬN TRỌNG)', level=1)
#             for w in warnings:
#                 warn_text = w if isinstance(w, str) else w.get('warning', w.get('description', ''))
#                 warn_desc = "" if isinstance(w, dict) else w.get('description', '')
                
#                 p = merged_doc.add_paragraph(style='List Bullet')
#                 r = p.add_run(warn_text)
#                 r.bold = True
#                 r.font.color.rgb = RGBColor(204, 153, 0) # Màu vàng cam
#                 if warn_desc:
#                     merged_doc.add_paragraph(warn_desc)
                    
#         if "legal_df" in analysis_dict and not analysis_dict["legal_df"].empty:
#             merged_doc.add_heading('3. RỦI RO PHÁP LÝ CHI TIẾT', level=1)
#             for _, row in analysis_dict["legal_df"].iterrows():
#                 p = merged_doc.add_paragraph(style='List Bullet')
#                 r = p.add_run(f"[{row.get('Mức độ', 'Vàng')}] {row.get('Vấn đề', '')}: ")
#                 r.bold = True
#                 r.font.color.rgb = RGBColor(204, 0, 0) # Màu đỏ
#                 p.add_run(str(row.get('Rủi ro', '')))
                
#                 p_sug = merged_doc.add_paragraph(f"Đề xuất: {row.get('Đề xuất', '')}")
#                 p_sug.runs[0].italic = True
        
#         merged_doc.add_page_break()
        
#         # --- PHẦN 2: CHÉP LẠI NỘI DUNG HỢP ĐỒNG GỐC ---
#         merged_doc.add_heading('--- NỘI DUNG HỢP ĐỒNG GỐC ---', 0)
#         for element in doc.element.body:
#             merged_doc.element.body.append(element)
            
#         # Lưu file xuất ra memory buffer
#         out_buffer = io.BytesIO()
#         merged_doc.save(out_buffer)
#         out_buffer.seek(0)
#         return out_buffer
#     except Exception as e:
#         print(f"Lỗi khi xử lý DOCX: {e}")
#         return None
    



# def generate_annotated_docx(analysis_dict: dict) -> io.BytesIO:
#     """Tạo một file DOCX chứa các comment và phân tích rủi ro."""
#     if not HAS_DOCX:
#         return None
        
#     doc = Document()
#     doc.add_heading('BÁO CÁO RÀ SOÁT HỢP ĐỒNG - GREENNODE GUARDIAN', 0)
    
#     # Tóm tắt
#     doc.add_heading('1. TÓM TẮT CHUNG', level=1)
#     doc.add_paragraph(analysis_dict.get("summary", ""))
    
#     # Cảnh báo
#     warnings = analysis_dict.get("yellow_warnings", [])
#     if warnings:
#         doc.add_heading('2. CẢNH BÁO (THẬN TRỌNG)', level=1)
#         for w in warnings:
#             warn_text = w if isinstance(w, str) else w.get('warning', w.get('description', ''))
#             warn_desc = "" if isinstance(w, dict) else w.get('description', '')
            
#             p = doc.add_paragraph(style='List Bullet')
#             r = p.add_run(warn_text)
#             r.bold = True
#             r.font.color.rgb = RGBColor(204, 153, 0) # Màu vàng cam
#             if warn_desc:
#                 doc.add_paragraph(warn_desc)
                
#     # Pháp lý
#     if "legal_df" in analysis_dict and not analysis_dict["legal_df"].empty:
#         doc.add_heading('3. RỦI RO PHÁP LÝ', level=1)
#         for _, row in analysis_dict["legal_df"].iterrows():
#             p = doc.add_paragraph(style='List Bullet')
#             r = p.add_run(f"[{row['Mức độ']}] {row['Vấn đề']}: ")
#             r.bold = True
#             r.font.color.rgb = RGBColor(204, 0, 0) # Màu đỏ
#             p.add_run(str(row['Rủi ro']))
            
#             p_sug = doc.add_paragraph(f"Đề xuất: {row['Đề xuất']}")
#             p_sug.runs[0].italic = True
            
#     # Lưu file vào memory
#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer


st.set_page_config(page_title="GreenNode Contract Guardian", layout="wide")
st.title("🛡️ GreenNode Contract Guardian")

@st.cache_data(show_spinner=False)
def cached_extract_risks(contract_text: str):
    return extract_contract_risks(contract_text)

@st.cache_data(show_spinner=False)
def cached_compliance_check(company_name: str):
    return check_compliance_mock(company_name)

@st.cache_data(show_spinner=False)
def cached_agent_response(contract_text: str, question: str):
    return run_agent(contract_text, question)

if "contract_text" not in st.session_state:
    st.session_state.contract_text = ""
if "session_id" not in st.session_state:
    st.session_state.session_id = str(uuid.uuid4())[:8]
if "analysis" not in st.session_state:
    st.session_state.analysis = None
if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

init_session(st.session_state.session_id)
memory = get_memory_client()

uploaded_file = st.file_uploader("Tải hợp đồng lên", type=["docx", "pdf", "txt"])
#model_name = st.selectbox("Chọn Model", ["google/gemma-4-31b-it", "qwen/qwen3-5-27b", "minimax/minimax-m2.5"], index=0)

# if uploaded_file is not None:
#     try:
#         st.session_state.contract_text = extract_text_from_file(uploaded_file.getvalue(), uploaded_file.name)
#         memory.add_event("system", f"Contract uploaded: {uploaded_file.name}")
#     except Exception as exc:
#         st.error(f"Không đọc được file: {exc}")

if uploaded_file is not None:
    try:
        file_bytes = uploaded_file.getvalue()
        # Lưu bytes gốc để sau này dùng cho xuất file DOCX
        st.session_state.original_file_bytes = file_bytes
        st.session_state.uploaded_filename = uploaded_file.name
        
        st.session_state.contract_text = extract_text_from_file(file_bytes, uploaded_file.name)
        memory.add_event("system", f"Contract uploaded: {uploaded_file.name}")
    except Exception as exc:
        st.error(f"Không đọc được file: {exc}")

if st.button("Phân tích hợp đồng", type="primary"):
    contract_text = st.session_state.contract_text.strip()
    if not contract_text:
        st.warning("Vui lòng tải lên hợp đồng trước.")
    else:
        memory.add_event("user", "Requested contract analysis")
        
        with st.spinner("🔍 Đang rà soát hợp đồng..."):
            risks_raw = cached_extract_risks(contract_text)
        
        with st.spinner("✓ Đang kiểm tra danh sách trừng phạt & tuân thủ..."):
            compliance_raw = cached_compliance_check("GreenNode Contract Counterparty")
        
        def parse_llm_json(raw_text):
            try:
                return json.loads(raw_text)
            except Exception:
                pass
            match = re.search(r'```(?:json)?\s*(\{.*?\})\s*```', raw_text, re.DOTALL)
            if match:
                try:
                    return json.loads(match.group(1))
                except Exception:
                    pass
            start = raw_text.find('{')
            end = raw_text.rfind('}')
            if start != -1 and end != -1:
                try:
                    return json.loads(raw_text[start:end+1])
                except Exception:
                    pass
            return None

        parsed_risks = parse_llm_json(risks_raw)
        
        if parsed_risks:
            risks = parsed_risks
        else:
            # Xóa sạch các thẻ markdown rác nếu bị rơi vào fallback
            clean_summary = risks_raw.replace("```json", "").replace("```", "").strip()
            risks = {"summary": clean_summary, "legal_risks": [], "financial_risks": [], "yellow_warnings": [], "negotiation_points": []}
            
        memory.add_event("assistant", f"Analysis complete: {risks.get('summary', '')[:100]}")



        # BẪY LỖI: Kiểm tra kiểu dữ liệu trước khi đẩy vào DataFrame
        legal_rows = []
        raw_legal = risks.get("legal_risks", [])
        
        # Đảm bảo raw_legal luôn là một list (tránh lỗi nếu LLM trả về chuỗi đơn)
        if not isinstance(raw_legal, list):
            raw_legal = [raw_legal]
            
        for item in raw_legal:
            if isinstance(item, dict):
                legal_rows.append({
                    "Nhóm điều khoản": item.get("group", "Chung"),
                    "Mức độ": item.get("severity", "Vàng"),
                    "Vấn đề": item.get("issue", item.get("risk", "")),
                    "Rủi ro": item.get("description", ""),
                    "Đề xuất": item.get("suggestion", ""),
                })
            else:
                # Fallback nếu item LLM trả về là String
                legal_rows.append({
                    "Nhóm điều khoản": "Chung",
                    "Mức độ": "Vàng",
                    "Vấn đề": "Cần lưu ý",
                    "Rủi ro": str(item),
                    "Đề xuất": "Rà soát lại điều khoản này",
                })



        # legal_rows = []
        # for item in (risks.get("legal_risks") or []):
        #     legal_rows.append({
        #         "Nhóm điều khoản": item.get("group", "Chung"),
        #         "Mức độ": item.get("severity", "Vàng"),
        #         "Vấn đề": item.get("issue", item.get("risk", "")),
        #         "Rủi ro": item.get("description", ""),
        #         "Đề xuất": item.get("suggestion", ""),
        #     })
        if not legal_rows:
            legal_rows = [{"Nhóm điều khoản": "Tổng quan", "Mức độ": "Vàng", "Vấn đề": "Chưa có phân tách chi tiết", "Rủi ro": "Cần rà soát thêm", "Đề xuất": "Dùng tool phân tích để bóc tách điều khoản"}]

        finance_df = pd.DataFrame([
            {"Chỉ số": "Tổng giá trị", "Giá trị": risks.get("total_value", "N/A")},
            {"Chỉ số": "Working Capital Gap", "Giá trị": risks.get("working_capital_gap", "N/A")},
            {"Chỉ số": "DSO ước tính", "Giá trị": risks.get("dso_estimate", "N/A")},
        ])

        cashflow_df = pd.DataFrame(risks.get("cashflow_table", [{"Kỳ": "Đợt 1", "Giá trị": "Chưa xác định", "Ghi chú": "Thiếu điều khoản thanh toán"}]))

        st.session_state.analysis = {
            "summary": risks.get("summary", "Đã phân tích hợp đồng."),
            "yellow_warnings": risks.get("yellow_warnings", []),
            "legal_df": pd.DataFrame(legal_rows),
            "finance_df": finance_df,
            "cashflow_df": cashflow_df,
            "compliance": compliance_raw,
            "email": generate_negotiation_email(json.dumps(risks.get("negotiation_points", []), ensure_ascii=False)),
        }

if st.session_state.analysis:
    st.divider()
    tab_summary, tab_legal, tab_finance, tab_compliance = st.tabs(["📑 Tóm tắt", "⚖️ Pháp lý", "💰 Tài chính", "🔎 Thẩm định & Đàm phán"])

    with tab_summary:
        summary_html = format_summary_html(
            st.session_state.analysis["summary"],
            st.session_state.analysis["yellow_warnings"]
        )
        st.markdown(summary_html, unsafe_allow_html=True)
        # Nút Download DOCX có comment được đặt ở đây theo đúng thiết kế UX
        # if HAS_DOCX:
        #     docx_buffer = generate_annotated_docx(st.session_state.analysis)
        #     if docx_buffer:
        #         st.download_button(
        #             label="📥 Tải DOCX có comment",
        #             data=docx_buffer,
        #             file_name="HopDong_Review_Co_Comment.docx",
        #             mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        #             type="primary"
        #       ) 

        # Nút Download DOCX có comment được đặt ở đây theo đúng thiết kế UX
        if HAS_DOCX and hasattr(st.session_state, 'original_file_bytes') and st.session_state.uploaded_filename.endswith('.docx'):
            docx_buffer = generate_annotated_docx(
                st.session_state.analysis, 
                st.session_state.original_file_bytes
            )
            if docx_buffer:
                st.download_button(
                    label="📥 Tải DOCX đã Review",
                    data=docx_buffer,
                    file_name=f"Reviewed_{st.session_state.uploaded_filename}",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

    with tab_legal:
        st.dataframe(st.session_state.analysis["legal_df"], width='stretch', hide_index=True)

    with tab_finance:
        metrics = st.session_state.analysis["finance_df"]
        
        # Dùng HTML/CSS Grid để tạo Metric Card nhỏ gọn và tự động responsive
        metric_html = """
<style>
.finance-metrics-container {
    display: grid;
    grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
    gap: 15px;
    margin-bottom: 25px;
}
.f-card {
    background: #ffffff;
    border: 1px solid #e2e8f0;
    border-radius: 8px;
    padding: 15px;
    box-shadow: 0 1px 3px rgba(0,0,0,0.05);
}
.f-title {
    font-size: 0.85rem;
    color: #64748b;
    font-weight: 600;
    margin-bottom: 8px;
    text-transform: uppercase;
}
.f-value {
    font-size: 1.15rem; /* Size nhỏ gọn hơn st.metric mặc định */
    color: #0f172a;
    font-weight: 700;
    word-wrap: break-word;
}
</style>
<div class="finance-metrics-container">
"""
        
        for idx in range(len(metrics)):
            title = metrics.iloc[idx]["Chỉ số"]
            val = metrics.iloc[idx]["Giá trị"]
            metric_html += f"""
<div class="f-card">
    <div class="f-title">{title}</div>
    <div class="f-value">{val}</div>
</div>
"""
            
        metric_html += "</div>"
        
        st.markdown(metric_html, unsafe_allow_html=True)
        
        st.markdown("### Bảng Dòng Tiền Dự Kiến")
        st.dataframe(st.session_state.analysis["cashflow_df"], width='stretch', hide_index=True)

    # with tab_compliance:
    #     try:
    #         compliance = json.loads(st.session_state.analysis["compliance"])
    #         compliance_html = format_compliance_html(compliance)
    #         st.markdown(compliance_html, unsafe_allow_html=True)
    #     except Exception:
    #         st.markdown(f"```json\n{st.session_state.analysis['compliance']}\n```")


    with tab_compliance:
        st.subheader("🔍 Kết Quả Thẩm Định (Due Diligence)")
        compliance_raw_str = str(st.session_state.analysis.get("compliance", ""))
        
        try:
            match = re.search(r'(\{.*\})', compliance_raw_str, re.DOTALL)
            if match:
                compliance = json.loads(match.group(1))
            else:
                compliance = json.loads(compliance_raw_str)
            
            # Xử lý trường hợp chuỗi bị stringify 2 lần
            if isinstance(compliance, str):
                compliance = json.loads(compliance)
                
            # Đã chuyển sang dùng UI mặc định của Streamlit (bỏ HTML thô)
            company_name = compliance.get("company_name", "Đối tác chưa định danh")
            status = compliance.get("partner_status", "Chưa xác định")
            
            st.info(f"**🏢 Thông tin Đối tác:** {company_name}\n\n**Trạng thái rủi ro:** {status}")
            
            col1, col2 = st.columns(2)
            with col1:
                st.warning(f"**🛡️ Trừng phạt & Danh sách đen (PEPs):**\n\n{compliance.get('ky_luat_peps', 'Không có thông tin')}")
            with col2:
                st.warning(f"**⚖️ Xung đột lợi ích:**\n\n{compliance.get('conflict_of_interest', 'Không có thông tin')}")
                
            hist = compliance.get("history", {})
            if isinstance(hist, dict):
                st.success(f"**📊 Lịch sử hợp tác:** Hợp đồng trước đây: `{hist.get('previous_contracts', 0)}` | Sự cố: `{hist.get('incidents', 0)}`")
                
            st.success(f"**💡 Đề xuất hành động:**\n\n{compliance.get('recommendation', 'Cần rà soát kỹ trước khi phê duyệt.')}")

        except Exception as e:
            # Fallback nếu LLM trả về chuỗi text bình thường không thể parse
            clean_text = compliance_raw_str.replace("```json", "").replace("```", "").strip()
            st.info(f"**📝 Báo cáo Thẩm định:**\n\n{clean_text}")
            
        st.markdown("---")
        st.markdown("### ✉️ Gợi ý Email Đàm Phán")
        email_content = st.session_state.analysis.get("email", "")
        if email_content:
            clean_email = email_content.replace("```markdown", "").replace("```", "").strip()
            st.info("Bản thảo email dưới đây được AI tổng hợp dựa trên các rủi ro đã phát hiện. Bạn có thể copy và điều chỉnh lại trước khi gửi cho đối tác.")
            st.text_area("Nội dung Email (Có thể chỉnh sửa):", value=clean_email, height=300)
        else:
            st.warning("Không có dữ liệu để tạo email gợi ý.")
    # with tab_compliance:
    #     compliance_raw_str = str(st.session_state.analysis.get("compliance", ""))
    #     try:
    #         match = re.search(r'(\{.*\})', compliance_raw_str, re.DOTALL)
    #         if match:
    #             compliance = json.loads(match.group(1))
    #         else:
    #             compliance = json.loads(compliance_raw_str)
                
    #         compliance_html = format_compliance_html(compliance)
    #         st.markdown(compliance_html, unsafe_allow_html=True)
    #         print("Compliance data parsed and displayed successfully. {}".format(compliance))
    #     except Exception:
    #         clean_text = compliance_raw_str.replace("```json", "").replace("```", "").strip()
    #         st.markdown(f"""
    #         <div style="background-color: #f8f9fa; padding: 20px; border-radius: 8px; border-left: 5px solid #6c757d; font-family: sans-serif; white-space: pre-wrap; line-height: 1.6; color: #333;">
    #             <strong style="color: #495057;">📝 Báo cáo Thẩm định:</strong><br><br>{clean_text}
    #         </div>
    #         """, unsafe_allow_html=True)
            
    #     st.markdown("---")
    #     st.markdown("### ✉️ Gợi ý Email Đàm Phán")
    #     email_content = st.session_state.analysis.get("email", "")
    #     if email_content:
    #         clean_email = email_content.replace("```markdown", "").replace("```", "").strip()
    #         st.info("Bản thảo email dưới đây được AI tổng hợp dựa trên các rủi ro đã phát hiện. Bạn có thể copy và điều chỉnh lại trước khi gửi cho đối tác.")
    #         st.text_area("Nội dung Email (Có thể chỉnh sửa):", value=clean_email, height=300)
    #     else:
    #         st.warning("Không có dữ liệu để tạo email gợi ý.")

st.divider()
st.subheader("💬 Hỏi thêm về hợp đồng")
quick_questions = ["Điều khoản nào rủi ro nhất?", "Tóm tắt nghĩa vụ của mỗi bên", "Thời hạn hợp đồng và gia hạn"]

cols = st.columns(len(quick_questions))
for i, question in enumerate(quick_questions):
    if cols[i].button(question, width='stretch'):
        memory.add_event("user", question)
        with st.spinner(f"⏳ Answering: {question}..."):
            answer = cached_agent_response(st.session_state.contract_text, question)
        memory.add_event("assistant", answer[:100])
        st.session_state.chat_history.append((question, answer))

for user_q, bot_a in st.session_state.chat_history:
    st.chat_message("user").write(user_q)
    st.chat_message("assistant").write(bot_a)

user_input = st.chat_input("Nhập câu hỏi về hợp đồng (VD: Soạn email xin giảm mức phạt xuống 4%)...")
if user_input:
    memory.add_event("user", user_input)
    st.chat_message("user").write(user_input)
    
    with st.spinner("⏳ Khởi chạy Agent..."):
        answer = cached_agent_response(st.session_state.contract_text, user_input)
        
    st.chat_message("assistant").write(answer)
    memory.add_event("assistant", answer[:100])
    st.session_state.chat_history.append((user_input, answer))
    
    if len(memory.get_session_events()) % 6 == 0:
        patterns = memory.extract_contract_patterns()
        if patterns.get("common_risks"):
            st.toast(f"📚 Agent đã ghi nhớ: {len(patterns['common_risks'])} điểm rủi ro để hỗ trợ đàm phán.")

with st.sidebar:
    st.markdown("### 🧠 Agent Memory Session")
    st.caption(f"Session ID: {st.session_state.session_id}")
    patterns = memory.extract_contract_patterns()
    st.json({"patterns": patterns, "interactions": len(memory.get_session_events())})